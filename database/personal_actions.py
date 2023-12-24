from aiogram import Bot, Dispatcher, executor, types
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.dispatcher import FSMContext
from database import keyboards as kb
from database import models as md
import db
from config import ADMIN_ID, BOT_TOKEN
import pandas as pd
from docx import Document
import os

storage = MemoryStorage()
bot = Bot(BOT_TOKEN)
dp = Dispatcher(bot=bot, storage=storage)


class NewOrder(StatesGroup):
    coming = State()
    price = State()


class ByTicket(StatesGroup):
    date_travel = State()
    place = State()


class DeleteBy(StatesGroup):
    id = State()

# Стартовая функция
@dp.message_handler(commands=['start'])
async def cmd_start(message: types.Message):
    # Добавление пользователя в БД
    user = md.Users.get_or_none(user_id = message.from_user.id)
    if not user:
        user_new = md.Users.create(user_id = message.from_user.id, name = message.from_user.first_name)
        user_new.save()
    await message.answer(f'{message.from_user.first_name}, добро пожаловать в Белгородский автовокзал!',
                         reply_markup=kb.main)
    # Проверка является ли пользователь администратором
    if message.from_user.id == int(ADMIN_ID):
        await message.answer(f'Вы авторизовались как администратор!', reply_markup=kb.main_admin)


# Вызов функции, если пользователь выбрал дату рейса
@dp.callback_query_handler(lambda c: c.data.startswith('add_to_by_'))
async def add_by_history(call: types.CallbackQuery):
    split_data = call.data.split('_')
    print(split_data)  # Добавляем эту строку для отображения значений, полученных после разделения
    # Извлекаем нужные данные
    coming = str(split_data[3])
    schedule_id = int(split_data[4])
    price = int(split_data[5])
    data = str(split_data[6])
    bus_id = int(split_data[7])
    user_id = call.from_user.id
    connection = await db.create_db_connection()
    # Извлекаем name автобуса по его ID
    bus = md.Buses.get_or_none(id = bus_id)
    # Извлекаем ID покупки, если пользователь уже покупал билет на этот рейс
    by_ticket = md.History.get_or_none(user_id = user_id, bus_name = bus.coming, date_travel = data)
    # Проверка на наличие билета на этот рейс (у пользователя может быть куплен только один билет на определенный рейс)
    if by_ticket:
        await call.message.answer("У вас уже куплен билет на этот рейс")
    # Если у пользователя нет билета на этот рейс
    else:
        # Считаем место пользователя
        count_places = md.Schedule.get_or_none(id = schedule_id)
        place = 51 - count_places.places
        # Если его место получилось больше, чем количество мест в автобусе, то свободных мест нет
        if place <= 50:
            # Уменьшаем количество свободных мест на 1
            await connection.execute("UPDATE public.schedule SET places = places - 1 WHERE id = $1", schedule_id)
            # Добавляем данные в историю покупок
            history_new = md.History.create(user_id=user_id,
                                            date_travel=data,
                                            place=place,
                                            bus_name=bus.coming,
                                            price=price)
            history_new.save()

            # Выводим билет пользователю
            await call.message.answer(
                f"Ваш билет:\n"
                f"Прибытие: {coming} \n"
                f"Дата рейса: {data} \n"
                f"Ваше место: {place}\n"
                f"Цена: {price}")
        else:
            await call.message.answer("На эту дату нет свободных мест")


# Вызов функции, если пользователь выбрал рейс
@dp.callback_query_handler(lambda c: c.data.startswith('add_to_history_'))
async def add_to_data(call: types.CallbackQuery):
    split_data = call.data.split('_')
    print(split_data)  # Добавляем эту строку для отображения значений, полученных после разделения
    # Извлекаем нужные данные
    coming = str(split_data[3])
    bus_id = int(split_data[4])
    price = int(split_data[5])
    # Извлекаем расписание автобуса по его ID
    places_bus = md.Schedule.select().where(md.Schedule.bus_id == bus_id)
    # Выводим расписание с кнопкой выбора даты
    for place in places_bus:
        keyboard = types.InlineKeyboardMarkup().add(
            types.InlineKeyboardButton(text="Выбрать дату",
                                       callback_data=f"add_to_by_{coming}_"
                                                     f"{place.id}_{price}_"
                                                     f"{place.data}_"
                                                     f"{bus_id}"))
        await call.message.answer(f"Прибытие: {coming}, "
                                  f"Дата: {place.data}, "
                                  f"Свободных мест: {place.places}", reply_markup=keyboard)


# Функция вывода рейсов
@dp.message_handler(text='Рейсы')
async def handle_show_landscapes(message: types.Message):
    # Извлекаем все рейсы
    buses = md.Buses.select()
    # Выводим рейсы с кнопкой покупки билета
    for bus in buses:
        keyboard = types.InlineKeyboardMarkup().add(
            types.InlineKeyboardButton(text="Купить билет",
                                       callback_data=
                                       f"add_to_history_{str(bus.coming)}_"
                                       f"{int(bus.id)}_{int(bus.price)}"))
        await message.answer(f"Прибытие: {bus.coming}, "
                             f"Цена: {bus.price}", reply_markup=keyboard)


# Вывод истории покупок пользователя
@dp.message_handler(text='История покупок')
async def handle_show_landscapes(message: types.Message):
    # Извлечение истории покупок пользователя по его ID

    history = md.History.select().where(md.History.user_id == message.from_user.id)
    # Проверка на наличие покупок
    if history.count() > 0:
        for by in history:
            await message.answer(
                f"Прибытие: {by.bus_name}, "
                f"Дата покупки: {by.date_by}, "
                f"Дата поездки: {by.date_travel}, "
                f"Ваше место: {by.place}")
    else:
        await message.answer(f'Вы еще ничего не покупали')


# Вывод панели администратора
@dp.message_handler(text='Панель администратора')
async def contacts(message: types.Message):
    # Проверка на то, что пользователь является админом
    if message.from_user.id == int(ADMIN_ID):
        # Поменять кнопки на панели
        await message.answer(f'Вы вошли в админ-панель', reply_markup=kb.admin_panel)
    else:
        await message.reply('Вы не администратор.')


# Вывод панели главного меню
@dp.message_handler(text='В главное меню')
async def process_main_menu(message: types.Message):
    # Проверка на то, что пользователь является администратором
    if message.from_user.id == int(ADMIN_ID):
        # Вывод панели для администратора
        await message.answer('Выберите действие:', reply_markup=kb.main_admin)
    else:
        # Вывод панели для пользователя
        await message.answer('Выберите действие:', reply_markup=kb.main)


# Вывод истории покупок всех пользователей
@dp.message_handler(text='История покупок всех пользователей')
async def handle_show_landscapes(message: types.Message):
    # Проверка на то, является ли пользователь админом
    if message.from_user.id == int(ADMIN_ID):
        # Извлечение покупок всех пользователей
        history = md.History.select()
        # Проверка на наличие покупок
        if history.count() > 0:
            for by in history:
                await message.answer(
                    f"ID: {by.id}, "
                    f"ID пользователя: {by.user_id}, "
                    f"Прибытие: {by.bus_name}, "
                    f"Дата покупки: {by.date_by}, "
                    f"Дата поездки: {by.date_travel}, "
                    f"Место: {by.place}")
        else:
            await message.answer(f'Еще ничего не покупали')
    else:
        await message.reply('Вы не администратор.')


@dp.message_handler(text='Контакты')
async def contacts(message: types.Message):
    await message.answer(f'Админ: @tejietty3uk')


# Удаление автобуса
@dp.message_handler(text='Удалить рейс')
async def delete_by_id(message: types.Message):
    if message.from_user.id == int(ADMIN_ID):
        await message.answer('Введите прибытие автобуса, который необходимо удалить')
        await DeleteBy.id.set()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')


# Удаление автобуса
@dp.message_handler(state=DeleteBy.id)
async def process_delete_by_id(message: types.Message, state: FSMContext):
    # Проверка на то, что пользователь является админом
    if message.from_user.id == int(ADMIN_ID):
        name = message.text
        bus = md.Buses.get_or_none(coming = name)
        # Проверка есть ли автобус с таким ID
        if bus:
            md.Schedule.delete().where(md.Schedule.bus_id == bus.id)
            bus.delete_instance()

            await message.answer(f'Автобус с прибытием в г. {name} успешно удален.')
        else:
            await message.answer('Рейс с таким прибытием не найден.')
            await state.finish()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')


# Добавление автобуса
@dp.message_handler(text='Добавить рейс')
async def add_item(message: types.Message):
    # Проверка на то, что пользователь является админом
    if message.from_user.id == int(ADMIN_ID):
        await NewOrder.coming.set()
        await message.answer(f'Введите место прибытия автобуса:')
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')


# Добавление автобуса
@dp.message_handler(state=NewOrder.coming)
async def add_item_type(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['coming'] = message.text
    await message.answer('Напишите цену билета:')
    await NewOrder.next()


# Добавление автобуса
@dp.message_handler(state=NewOrder.price)
async def add_item_type(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['price'] = int(message.text)

        new = md.Buses(coming = data['coming'], price= data['price'])
        new.save()

    connection = await db.create_db_connection()
    # Извлекаем ID нового автобуса
    bus_id = md.Buses.get_or_none(coming = data['coming'])
    # Создаем рейсы на каждую дату
    for i in range(1, 11):
        # Заполняем дату нужного рейса
        if i < 10:
            date = md.Schedule(data = f"2024/01/0{i}", bus_id = bus_id)
        else:
            date = md.Schedule(data=f"2024/01/{i}", bus_id=bus_id)
        date.save()

    await message.answer('Товар успешно создан!')
    await state.finish()


@dp.message_handler(text='Сделать выгрузку данных')
async def import_to(message: types.Message):
    if message.from_user.id == int(ADMIN_ID):
        # Удаление файла, если уже существует
        if os.path.exists('by_history.xlsx'):
            os.remove('by_history.xlsx')
        if os.path.exists('by_history.docx'):
            os.remove('by_history.docx')

        # Выборка данных из таблицы History
        query = md.History.select()

        # Преобразование данных из выборки в DataFrame
        history_data = list(query.dicts())
        df = pd.DataFrame(history_data)

        # Экспорт DataFrame в Excel файл
        df.to_excel('by_history.xlsx', index=False)

        # Создание файла docx и запись данных из DataFrame построчно
        doc = Document()
        doc.add_heading('История Заказов', 0)

        for index, row in df.iterrows():
            for column in df.columns:
                doc.add_paragraph(f"{column}: {row[column]}")
            doc.add_paragraph('')

        doc.save('by_history.docx')

        # Отправка файлов пользователю
        with open('by_history.xlsx', 'rb') as file_xlsx, open('by_history.docx', 'rb') as file_docx:
            await message.reply_document(file_xlsx, caption="Excel файл")
            await message.reply_document(file_docx, caption="Word файл")

    else:
        await message.reply('У Вас недостаточно прав')

# Если пользователь ввел не команду
@dp.message_handler()
async def answer(message: types.Message):
    await message.reply('Я тебя не понимаю')
